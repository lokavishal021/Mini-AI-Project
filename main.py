from flask import Flask, request, jsonify, render_template, session
import datetime
import re
import json
import os
from sympy import symbols, Eq, solve, simplify, diff, integrate, limit, sympify
import random

from docx import Document

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'a-strong-default-secret-key-for-dev')  # Use environment variable

# --- Load Code Snippets from JSON (with error handling) ---
try:
    if os.path.exists('programs.json'):
        with open('programs.json') as f:
            code_snippets = json.load(f)
    else:
        print("Warning: programs.json not found. Creating empty dict.")
        code_snippets = {}
except Exception as e:
    print(f"Error loading programs.json: {e}")
    code_snippets = {}

# --- Image Captioning Setup (lazy loading) ---
processor = None
model = None

def load_image_models():
    """Lazy load the image captioning models only when needed"""
    global processor, model
    if processor is None or model is None:
        try:
            from transformers import BlipProcessor, BlipForConditionalGeneration
            print("Loading BLIP models... This may take a while.")
            processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
            model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
            print("Models loaded successfully!")
        except Exception as e:
            print(f"Error loading models: {e}")
            raise

def get_image_caption(image):
    """Get caption for an image"""
    try:
        load_image_models()
        inputs = processor(images=image, return_tensors="pt")
        out = model.generate(**inputs)
        caption = processor.decode(out[0], skip_special_tokens=True)
        return caption
    except Exception as e:
        return f"Error generating caption: {str(e)}"

# --- Document Summarization ---
def summarize_text(text, max_sentences=5):
    """Extract main points from text"""
    try:
        # Split into sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        # Simple extractive summarization - take first few sentences and some from middle/end
        if len(sentences) <= max_sentences:
            summary = ". ".join(sentences) + "."
        else:
            # Take sentences from different parts of the document
            indices = [0, 1]  # First two sentences
            middle_idx = len(sentences) // 2
            indices.append(middle_idx)
            
            # Add more if needed
            remaining = max_sentences - len(indices)
            step = len(sentences) // (remaining + 1)
            for i in range(remaining):
                idx = step * (i + 1)
                if idx < len(sentences) and idx not in indices:
                    indices.append(idx)
            
            indices.sort()
            selected_sentences = [sentences[i] for i in indices if i < len(sentences)]
            summary = ". ".join(selected_sentences) + "."
        
        # Add bullet points for main points
        main_points = []
        for sentence in sentences[:10]:  # Look at first 10 sentences for keywords
            # Extract potential key points (sentences with important keywords)
            if any(word in sentence.lower() for word in ['important', 'main', 'key', 'significant', 'crucial', 'essential', 'primary']):
                main_points.append(sentence.strip())
        
        result = "📄 **Summary:**\n\n" + summary
        
        if main_points:
            result += "\n\n**Main Points:**\n"
            for i, point in enumerate(main_points[:3], 1):
                result += f"\n{i}. {point}"
        
        return result
    except Exception as e:
        return f"Error summarizing text: {str(e)}"

# --- Story Generation --- 
def generate_story(key_points):
    story_templates = [
        "Once upon a time, there was a {character} who was {adjective}. One day, they discovered {discovery}. This discovery led them on a journey to {destination}, where they encountered {obstacle}. With determination and courage, they overcame the obstacle and {resolution}.",
        "In a distant land, a {character} set out on an adventure to {goal}. Along the way, they faced many challenges, including {challenge}. But through wisdom and bravery, they succeeded in {achievement}. Their journey became a legend, known far and wide as the {story_name}.",
        "A {adjective} {character} found themselves caught in an unexpected situation. While trying to {action}, they stumbled upon {discovery}. This started a chain of events that led to {unexpected_turn}. In the end, they learned {lesson}, and their life was changed forever."
    ]

    story_data = {
        "character": "knight" if "knight" in str(key_points) else "hero",
        "adjective": "brave" if "brave" in str(key_points) else "kind",
        "discovery": "a hidden treasure" if "treasure" in str(key_points) else "a powerful artifact",
        "destination": "a distant castle" if "castle" in str(key_points) else "an enchanted forest",
        "obstacle": "a dangerous dragon" if "dragon" in str(key_points) else "an evil sorcerer",
        "resolution": "became a legend" if "legend" in str(key_points) else "defeated the dark forces",
        "goal": "defeat the evil forces" if "evil" in str(key_points) else "find a rare artifact",
        "challenge": "treacherous terrain" if "terrain" in str(key_points) else "a fierce monster",
        "achievement": "saving the kingdom" if "kingdom" in str(key_points) else "finding the treasure",
        "story_name": "The Brave Knight's Quest" if "knight" in str(key_points) else "The Hero's Journey",
        "action": "fight the sorcerer" if "fight" in str(key_points) else "seek the hidden treasure",
        "unexpected_turn": "they realized the treasure was cursed" if "cursed" in str(key_points) else "they were betrayed by an ally",
        "lesson": "the true meaning of courage" if "courage" in str(key_points) else "the importance of friendship"
    }

    story_template = random.choice(story_templates)
    story = story_template.format(**story_data)
    return story

# --- Basic Math Expression Evaluation ---
def evaluate_math_expression(expression):
    try:
        expression = expression.lower()
        expression = expression.replace("plus", "+").replace("minus", "-")
        expression = expression.replace("x", "*").replace("into", "*")
        expression = expression.replace("divided by", "/").replace("mod", "%")
        if re.match(r"^[\d\s\+\-\*/%\.\(\)]+$", expression):
            result = eval(expression)
            return f"The answer is: {result}"
        return None
    except Exception:
        return None

# --- Advanced Math Solver ---
def advanced_math_solver(expression):
    try:
        x = symbols('x')
        expression = expression.lower().replace("^", "**")

        if "solve" in expression:
            eq = expression.replace("solve", "").strip()
            lhs, rhs = eq.split("=")
            equation = Eq(sympify(lhs), sympify(rhs))
            result = solve(equation, x)
            return f"Solution: {result}"
        elif "differentiate" in expression or "derivative" in expression:
            expr = expression.replace("differentiate", "").replace("derivative", "").strip()
            return f"Derivative: {diff(sympify(expr))}"
        elif "integrate" in expression:
            expr = expression.replace("integrate", "").strip()
            return f"Integral: {integrate(sympify(expr))}"
        elif "simplify" in expression:
            expr = expression.replace("simplify", "").strip()
            return f"Simplified: {simplify(sympify(expr))}"
        elif "limit" in expression:
            expr = expression.replace("limit", "").strip()
            return f"Limit as x approaches ∞: {limit(sympify(expr), x, float('inf'))}"
        return None
    except Exception as e:
        return f"Sorry, I couldn't solve that. Error: {str(e)}"

# --- Disease Info Handler ---
def get_disease_info(message):
    disease_data = {
        "cold": {
            "name": "Common Cold",
            "symptoms": "Runny nose, sore throat, cough, congestion, slight body aches.",
            "cause": "Caused by a viral infection (usually rhinovirus).",
            "treatment": "Rest, hydration, and over-the-counter cold medications.",
            "severity": "Mild"
        },
        "fever": {
            "name": "Fever",
            "symptoms": "High body temperature, chills, sweating, headache, body aches.",
            "cause": "Usually due to an infection (bacterial or viral).",
            "treatment": "Stay hydrated, take paracetamol or ibuprofen, and rest.",
            "severity": "Mild to Moderate"
        },
        "covid": {
            "name": "COVID-19",
            "symptoms": "Fever, cough, fatigue, shortness of breath, loss of taste or smell.",
            "cause": "Caused by SARS-CoV-2 virus, spreads through droplets.",
            "treatment": "Isolation, monitoring symptoms, and seeking medical help if needed.",
            "severity": "Varies from Mild to Severe"
        },
        "malaria": {
            "name": "Malaria",
            "symptoms": "Fever, chills, vomiting, headache, muscle pain.",
            "cause": "Spread by Anopheles mosquitoes carrying Plasmodium parasite.",
            "treatment": "Antimalarial medications prescribed by doctors.",
            "severity": "Moderate to Severe"
        },
        "diabetes": {
            "name": "Diabetes",
            "symptoms": "Increased thirst, frequent urination, fatigue, blurred vision.",
            "cause": "High blood sugar due to insulin issues (Type 1 or 2).",
            "treatment": "Managed with medication, insulin, diet control, and exercise.",
            "severity": "Chronic"
        },
        "hypertension": {
            "name": "Hypertension",
            "symptoms": "Often silent, may include headache, shortness of breath, or nosebleeds.",
            "cause": "High pressure in the arteries. Risk factor for heart disease.",
            "treatment": "Lifestyle changes and antihypertensive drugs.",
            "severity": "Chronic"
        },
        "headache": {
            "name": "Headache",
            "symptoms": "Pain in head, scalp, or neck. Can be dull or sharp.",
            "cause": "Stress, dehydration, sinus issues, eye strain, or more serious causes.",
            "treatment": "Rest, hydration, and over-the-counter pain relievers.",
            "severity": "Mild to Moderate"
        }
    }

    for keyword, info in disease_data.items():
        if keyword in message:
            return (
                f"🩺 *{info['name']}*\n"
                f"- **Symptoms**: {info['symptoms']}\n"
                f"- **Cause**: {info['cause']}\n"
                f"- **Treatment**: {info['treatment']}\n"
                f"- **Severity**: {info['severity']}"
            )
    return None

# --- Program Snippet Handler ---
def get_program_snippet(message):
    for key in code_snippets:
        if key.lower() in message.lower():
            return f"Here is the {key} program:\n```python\n{code_snippets[key]}\n```"
    return None

# --- Wikipedia Info ---
def get_wikipedia_info(query, more=False):
    try:
        import wikipedia

        if not more:
            session['wiki_topic'] = query
            session['wiki_offset'] = 0
        else:
            query = session.get('wiki_topic')
            session['wiki_offset'] = session.get('wiki_offset', 0) + 2

        if not query:
            return "Please ask about a topic first."
        
        offset = session.get('wiki_offset', 0)
        summary = wikipedia.summary(query, sentences=offset + 2)
        paragraphs = summary.split(". ")
        start = offset
        end = start + 2
        more_info = ". ".join(paragraphs[start:end])
        return more_info if more_info else "No more information available."
    except Exception as e:
        return f"Sorry, I couldn't find information on that topic. Error: {str(e)}"

# --- Current Affairs ---
def get_current_affairs():
    """Fetches top news headlines from NewsAPI."""
    try:
        import requests

        # API Key for NewsAPI.org. For better security, consider moving this to an environment variable.

        api_key = "26d03ad1d52f4222a123717cc2dea0b5"
        
        # Fetch top headlines from India. You can change the country code (e.g., 'us' for USA).
        url = f"https://newsapi.org/v2/top-headlines?country=in&apiKey={api_key}"
        response = requests.get(url)
        response.raise_for_status()  # Raise an exception for bad status codes
        
        articles = response.json().get("articles", [])
        if not articles:
            return "Sorry, I couldn't fetch the latest news right now."

        headlines = ["📰 Here are the top 5 headlines:"] + [f"- {article['title']}" for article in articles[:5]]
        return "\n".join(headlines)
    except Exception as e:
        print(f"Error fetching current affairs: {e}")
        return "Sorry, I'm having trouble fetching the news at the moment."

# --- Assistant Logic ---
def assistant_logic(send):
    data_btn = send.lower()

    # Check if user is responding to document action request
    if 'pending_document' in session and session['pending_document']:
        if "read" in data_btn or "read out" in data_btn or "full" in data_btn:
            content = session.get('document_content', '')
            session['pending_document'] = False
            return {
                "text": "📖 Reading the full document...",
                "speak": content,
                "full_content": content
            }
        elif "summarize" in data_btn or "summary" in data_btn or "short" in data_btn or "main points" in data_btn:
            content = session.get('document_content', '')
            session['pending_document'] = False
            summary = summarize_text(content)
            return {
                "text": summary,
                "speak": summary,
                "full_content": None
            }
        else:
            return {
                "text": "Please choose an option: 'read out' for full content or 'summarize' for main points.",
                "speak": "Please choose read out or summarize",
                "full_content": None
            }

    if "what is your name" in data_btn:
        return "My name is Virtual Assistant"
    elif any(greet in data_btn for greet in ["hello", "hye", "hay", "hi"]):
        return "Hey sir, how can I help you!"
    elif "how are you" in data_btn:
        return "I am doing great these days, sir."
    elif "thanku" in data_btn or "thank" in data_btn:
        return "It's my pleasure, sir, to stay with you."
    elif "good morning" in data_btn:
        return "Good morning sir, I think you might need some help."
    elif "time now" in data_btn:
        now = datetime.datetime.now()
        return now.strftime("Current time is %I:%M %p")
    elif "current affairs" in data_btn:
        return get_current_affairs()
    elif "open youtube" in data_btn:
        return "OPEN_YOUTUBE"
    elif "open google" in data_btn:
        return "OPEN_GOOGLE"
    elif "open facebook" in data_btn:
        return "OPEN_FACEBOOK"
    elif "open sbtet" in data_btn:
        return "OPEN_SBTET"
    elif "open music" in data_btn:
        return "OPEN_MUSIC"
    elif "shutdown" in data_btn or "quit" in data_btn:
        return "Ok sir. Shutting down."

    if data_btn.startswith(("about ", "who is ", "what is ")):
        topic = data_btn.replace("about", "").replace("who is", "").replace("what is", "").strip()
        return get_wikipedia_info(topic, more=False)
    elif "more about him" in data_btn or "more about her" in data_btn:
        return get_wikipedia_info("", more=True)

    program_result = get_program_snippet(data_btn)
    if program_result:
        return program_result

    basic_result = evaluate_math_expression(data_btn)
    if basic_result:
        return basic_result

    advanced_result = advanced_math_solver(data_btn)
    if advanced_result:
        return advanced_result

    disease_result = get_disease_info(data_btn)
    if disease_result:
        return disease_result

    if "tell me a story" in data_btn:
        key_points = {"character": "young prince", "setting": "magical forest", "conflict": "an evil dragon", "resolution": "outsmarting the dragon using clever tricks"}
        story = generate_story(key_points)
        return story

    return "Sorry, I didn't understand that. Try asking about diseases, math problems, or say 'open YouTube' or upload a file or image."

# --- File Reading Functions ---
def read_pdf_file(file):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def read_docx_file(file):
    """Read DOCX file with fixed import"""
    try:
        import io

        # Read file content into memory
        file_content = file.read()
        file_stream = io.BytesIO(file_content)
        
        # Open document from stream
        doc = Document(file_stream)
        
        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))

        result = "\n".join(full_text)

        if not result.strip():
            return "The document appears to be empty."
        
        return result
        
    except ImportError:
        return "Error: python-docx library not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"

def read_txt_file(file):
    try:
        return file.read().decode("utf-8")
    except Exception as e:
        return f"Error reading TXT: {str(e)}"

# --- Image Recognition and Captioning ---
@app.route('/upload', methods=['POST'])
def upload():
    file = request.files.get('file') or request.files.get('image')
    if file:
        filename = file.filename.lower()
        try:
            if filename.endswith(('.png', '.jpg', '.jpeg', '.bmp')): 
                from PIL import Image
                import io
                img = Image.open(io.BytesIO(file.read()))
                caption = get_image_caption(img)

                weapons = ["gun", "knife", "pistol", "bomb", "rifle"]
                detected_weapons = [weapon for weapon in weapons if weapon in caption.lower()]

                if detected_weapons:
                    warning_message = f"⚠️ Warning: Possible weapon detected ({', '.join(detected_weapons)})."
                    caption = warning_message + "\n" + caption

                return jsonify({"caption": caption})

            elif filename.endswith('.pdf'):
                content = read_pdf_file(file)
                # Store content in session for later processing
                session['document_content'] = content
                session['pending_document'] = True
                return jsonify({
                    "type": "document",
                    "message": "📄 Document uploaded successfully! What would you like me to do?\n\n1️⃣ Type 'read out' - I'll read the full document\n2️⃣ Type 'summarize' - I'll extract main points for better understanding",
                    "speak": "Document uploaded. What would you like me to do? Say read out for full content, or summarize for main points."
                })
            
            elif filename.endswith('.docx'):
                content = read_docx_file(file)
                
                # Check if there was an error reading the file
                if content.startswith("Error"):
                    return jsonify({"status": "error", "message": content})
                
                # Store content in session for later processing
                session['document_content'] = content
                session['pending_document'] = True
                return jsonify({
                    "type": "document",
                    "message": "📄 Document uploaded successfully! What would you like me to do?\n\n1️⃣ Type 'read out' - I'll read the full document\n2️⃣ Type 'summarize' - I'll extract main points for better understanding",
                    "speak": "Document uploaded. What would you like me to do? Say read out for full content, or summarize for main points."
                })
            
            elif filename.endswith('.txt'):
                content = read_txt_file(file)
                session['document_content'] = content
                session['pending_document'] = True
                return jsonify({
                    "type": "document",
                    "message": "📄 Text file uploaded successfully! What would you like me to do?\n\n1️⃣ Type 'read out' - I'll read the full document\n2️⃣ Type 'summarize' - I'll extract main points for better understanding",
                    "speak": "Text file uploaded. What would you like me to do? Say read out for full content, or summarize for main points."
                })
            else:
                return jsonify({"status": "error", "message": "Unsupported file type"})
        
        except Exception as e:
            return jsonify({"status": "error", "message": str(e)})
    
    return jsonify({"status": "no file uploaded"})

# --- Routes ---
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/chat", methods=["POST"])
def chat():
    user_message = request.json.get("message", "")
    reply = assistant_logic(user_message)
    
    # Handle different reply types
    if isinstance(reply, dict):
        return jsonify(reply)
    else:
        return jsonify({"reply": reply})

# Test route to verify server is running
@app.route("/test")
def test():
    return "Flask server is running!"

if __name__ == "__main__":
    print("=" * 50)
    print("Starting Flask Virtual Assistant...")
    print("=" * 50)
    print(f"Server will start at: http://127.0.0.1:5000/")
    print("Press CTRL+C to quit")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5000)